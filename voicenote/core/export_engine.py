"""
VoiceNote Export Engine.
Handles document generation and export of notes, AI summaries, tasks,
and transcripts to PDF, Word (DOCX), and Plain Text (TXT) formats.
"""

import os
from pathlib import Path
from datetime import datetime
from typing import Dict, Any, List, Optional, Union


class ExportEngine:
    """Engine for exporting VoiceNote content into PDF, DOCX, and TXT formats."""

    DEFAULT_OPTIONS = {
        "include_summary": True,
        "include_tasks": True,
        "include_transcript": True,
        "include_metadata": True,
    }

    def __init__(self):
        pass

    def export(
        self,
        format_type: str,
        note_data: Union[Dict[str, Any], Any],
        output_path: Optional[str] = None,
        options: Optional[Dict[str, bool]] = None,
    ) -> str:
        """Central export dispatcher.

        Args:
            format_type: 'pdf', 'docx', or 'txt' (case-insensitive)
            note_data: Dictionary or object with note details
            output_path: Destination file path. If None, generates a default
              path.
            options: Section inclusion flags.

        Returns:
            str: The absolute path of the generated export file.
        """
        normalized_data = self._normalize_note_data(note_data)
        merged_options = {**self.DEFAULT_OPTIONS, **(options or {})}
        fmt = format_type.strip().lower().replace(".", "")

        if not output_path:
            title_slug = "".join(
                c if c.isalnum() or c in ("-", "_") else "_"
                for c in normalized_data.get("title", "VoiceNote").replace(
                    " ", "_"
                )
            )[:40].strip("_")
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"{title_slug}_{timestamp}.{fmt}"
            output_path = str(Path.cwd() / filename)

        # Ensure directory exists
        dest_dir = Path(output_path).parent
        dest_dir.mkdir(parents=True, exist_ok=True)

        if fmt == "pdf":
            return self.export_pdf(
                normalized_data, str(output_path), merged_options
            )
        elif fmt == "docx":
            return self.export_docx(
                normalized_data, str(output_path), merged_options
            )
        elif fmt in ("txt", "text", "md"):
            return self.export_txt(
                normalized_data, str(output_path), merged_options
            )
        else:
            raise ValueError(
                f"Unsupported export format: '{format_type}'. Supported formats: 'pdf', 'docx', 'txt'."
            )

    def _normalize_note_data(
        self, note_data: Union[Dict[str, Any], Any]
    ) -> Dict[str, Any]:
        """Normalize raw dictionary or dataclass/model objects into a standard dictionary."""
        if hasattr(note_data, "__dict__"):
            data = dict(note_data.__dict__)
        elif isinstance(note_data, dict):
            data = dict(note_data)
        else:
            data = {"title": str(note_data)}

        # Extract title
        data.setdefault("title", "Untitled Voice Note")
        data.setdefault(
            "created_at", datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        )
        data.setdefault("duration", "00:00")
        data.setdefault("category", "General")

        # Tags / Topics
        tags = (
            data.get("main_topics")
            or data.get("tags")
            or data.get("category", [])
        )
        if isinstance(tags, str):
            tags = [t.strip() for t in tags.split(",") if t.strip()]
        elif not isinstance(tags, list):
            tags = [str(tags)]
        data["tags"] = tags

        # Summary
        summary = data.get("summary") or ""
        data["summary"] = summary

        # Key points
        key_points = data.get("key_points") or []
        if isinstance(key_points, str):
            key_points = [
                kp.strip()
                for kp in key_points.split("\n")
                if kp.strip() and not kp.startswith("#")
            ]
        data["key_points"] = key_points

        # Tasks
        tasks = data.get("tasks") or []
        norm_tasks = []
        for t in tasks:
            if isinstance(t, dict):
                norm_tasks.append(
                    {
                        "title": t.get("title")
                        or t.get("desc")
                        or t.get("description")
                        or "Task item",
                        "priority": t.get("priority", "Medium"),
                        "assignee": t.get("assignee", "Unassigned"),
                        "due_date": t.get("due_date", "TBD"),
                        "status": t.get("status")
                        or ("Completed" if t.get("done") else "Pending"),
                    }
                )
            elif hasattr(t, "title"):
                norm_tasks.append(
                    {
                        "title": getattr(t, "title", "Task item"),
                        "priority": getattr(t, "priority", "Medium"),
                        "assignee": getattr(t, "assignee", "Unassigned"),
                        "due_date": getattr(t, "due_date", "TBD"),
                        "status": getattr(t, "status", "Pending"),
                    }
                )
        data["tasks"] = norm_tasks

        # Transcript
        transcript = (
            data.get("cleaned_text")
            or data.get("raw_text")
            or data.get("transcript")
            or ""
        )
        if isinstance(transcript, list):
            transcript = "\n".join(str(item) for item in transcript)
        data["transcript"] = str(transcript)

        return data

    def export_pdf(
        self,
        note_data: Dict[str, Any],
        output_path: str,
        options: Dict[str, bool],
    ) -> str:
        """Export note to a professionally styled PDF document using ReportLab."""
        try:
            from reportlab.lib.pagesizes import letter
            from reportlab.lib import colors
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.platypus import (
                SimpleDocTemplate,
                Paragraph,
                Spacer,
                Table,
                TableStyle,
                HRFlowable,
                KeepTogether,
            )
            from reportlab.lib.units import inch
        except ImportError as e:
            raise ImportError(
                f"ReportLab is required for PDF export: {e}. Install with 'pip install reportlab'."
            )

        doc = SimpleDocTemplate(
            output_path,
            pagesize=letter,
            rightMargin=40,
            leftMargin=40,
            topMargin=40,
            bottomMargin=40,
        )

        styles = getSampleStyleSheet()

        # Custom Brand Color Palette
        c_primary = colors.HexColor("#1E2B4B")  # Deep Navy
        c_accent = colors.HexColor("#6D59A7")  # Warm Purple
        c_light = colors.HexColor("#F8F7F4")  # Light Cream
        c_border = colors.HexColor("#E2DDD3")  # Soft Border
        c_dark = colors.HexColor("#2D3748")  # Body Text
        c_sub = colors.HexColor("#718096")  # Subtitle
        c_badge_bg = colors.HexColor("#EDE8DF")  # Badge Background

        # Custom Typography Styles
        title_style = ParagraphStyle(
            "DocTitle",
            parent=styles["Normal"],
            fontName="Helvetica-Bold",
            fontSize=22,
            leading=26,
            textColor=c_primary,
            spaceAfter=4,
        )

        subtitle_style = ParagraphStyle(
            "DocSubtitle",
            parent=styles["Normal"],
            fontName="Helvetica",
            fontSize=10,
            leading=14,
            textColor=c_sub,
            spaceAfter=12,
        )

        h2_style = ParagraphStyle(
            "DocHeading2",
            parent=styles["Normal"],
            fontName="Helvetica-Bold",
            fontSize=13,
            leading=17,
            textColor=c_primary,
            spaceBefore=14,
            spaceAfter=6,
        )

        body_style = ParagraphStyle(
            "DocBody",
            parent=styles["Normal"],
            fontName="Helvetica",
            fontSize=10,
            leading=14,
            textColor=c_dark,
            spaceAfter=6,
        )

        callout_style = ParagraphStyle(
            "DocCallout",
            parent=styles["Normal"],
            fontName="Helvetica-Oblique",
            fontSize=10,
            leading=14,
            textColor=c_primary,
        )

        meta_label_style = ParagraphStyle(
            "MetaLabel",
            parent=styles["Normal"],
            fontName="Helvetica-Bold",
            fontSize=9,
            leading=12,
            textColor=c_primary,
        )

        meta_val_style = ParagraphStyle(
            "MetaVal",
            parent=styles["Normal"],
            fontName="Helvetica",
            fontSize=9,
            leading=12,
            textColor=c_dark,
        )

        table_header_style = ParagraphStyle(
            "TableHeader",
            parent=styles["Normal"],
            fontName="Helvetica-Bold",
            fontSize=9,
            leading=12,
            textColor=colors.white,
        )

        table_cell_style = ParagraphStyle(
            "TableCell",
            parent=styles["Normal"],
            fontName="Helvetica",
            fontSize=8.5,
            leading=11,
            textColor=c_dark,
        )

        story = []

        # 1. Header Section
        story.append(
            Paragraph("VOICENOTE AI • DOCUMENT EXPORT", subtitle_style)
        )
        story.append(Paragraph(note_data.get("title", "Voice Note"), title_style))
        story.append(
            HRFlowable(
                width="100%",
                thickness=1.5,
                color=c_accent,
                spaceBefore=2,
                spaceAfter=12,
            )
        )

        # 2. Metadata Box (if enabled)
        if options.get("include_metadata", True):
            tags_str = ", ".join(
                [f"#{t.lstrip('#')}" for t in note_data.get("tags", [])]
            ) or "General"
            meta_data = [
                [
                    Paragraph("<b>Created Date:</b>", meta_label_style),
                    Paragraph(str(note_data.get("created_at", "")), meta_val_style),
                    Paragraph("<b>Audio Duration:</b>", meta_label_style),
                    Paragraph(str(note_data.get("duration", "00:00")), meta_val_style),
                ],
                [
                    Paragraph("<b>Categories / Tags:</b>", meta_label_style),
                    Paragraph(tags_str, meta_val_style),
                    Paragraph("<b>Generated By:</b>", meta_label_style),
                    Paragraph("VoiceNote AI Desktop", meta_val_style),
                ],
            ]
            meta_table = Table(
                meta_data, colWidths=[1.3 * inch, 2.3 * inch, 1.3 * inch, 2.3 * inch]
            )
            meta_table.setStyle(
                TableStyle(
                    [
                        ("BACKGROUND", (0, 0), (-1, -1), c_light),
                        ("BOX", (0, 0), (-1, -1), 1, c_border),
                        ("INNERGRID", (0, 0), (-1, -1), 0.5, c_border),
                        ("TOPPADDING", (0, 0), (-1, -1), 5),
                        ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
                        ("LEFTPADDING", (0, 0), (-1, -1), 8),
                        ("RIGHTPADDING", (0, 0), (-1, -1), 8),
                    ]
                )
            )
            story.append(meta_table)
            story.append(Spacer(1, 10))

        # 3. AI Executive Summary (if enabled)
        if options.get("include_summary", True):
            summary_text = note_data.get("summary")
            key_points = note_data.get("key_points", [])

            if summary_text or key_points:
                story.append(Paragraph("AI Executive Summary", h2_style))
                if summary_text:
                    sum_box_data = [
                        [
                            Paragraph(
                                summary_text.replace("\n", "<br/>"),
                                callout_style,
                            )
                        ]
                    ]
                    sum_table = Table(sum_box_data, colWidths=[7.2 * inch])
                    sum_table.setStyle(
                        TableStyle(
                            [
                                ("BACKGROUND", (0, 0), (-1, -1), c_light),
                                ("BOX", (0, 0), (-1, -1), 1, c_accent),
                                ("LEFTPADDING", (0, 0), (-1, -1), 12),
                                ("RIGHTPADDING", (0, 0), (-1, -1), 12),
                                ("TOPPADDING", (0, 0), (-1, -1), 8),
                                ("BOTTOMPADDING", (0, 0), (-1, -1), 8),
                            ]
                        )
                    )
                    story.append(sum_table)
                    story.append(Spacer(1, 8))

                if key_points:
                    story.append(Paragraph("<b>Key Takeaways:</b>", body_style))
                    for kp in key_points:
                        story.append(
                            Paragraph(f"• {kp}", body_style)
                        )
                    story.append(Spacer(1, 8))

        # 4. Action Items & Extracted Tasks (if enabled)
        if options.get("include_tasks", True):
            tasks = note_data.get("tasks", [])
            if tasks:
                story.append(Paragraph("Extracted Action Items & Tasks", h2_style))
                task_rows = [
                    [
                        Paragraph("Task Description", table_header_style),
                        Paragraph("Priority", table_header_style),
                        Paragraph("Assignee", table_header_style),
                        Paragraph("Due Date", table_header_style),
                        Paragraph("Status", table_header_style),
                    ]
                ]

                for t in tasks:
                    status_text = t.get("status", "Pending")
                    priority_text = t.get("priority", "Medium")
                    task_rows.append(
                        [
                            Paragraph(t.get("title", ""), table_cell_style),
                            Paragraph(f"<b>{priority_text}</b>", table_cell_style),
                            Paragraph(t.get("assignee", "Unassigned"), table_cell_style),
                            Paragraph(t.get("due_date", "TBD"), table_cell_style),
                            Paragraph(status_text, table_cell_style),
                        ]
                    )

                task_table = Table(
                    task_rows,
                    colWidths=[3.2 * inch, 0.9 * inch, 1.2 * inch, 1.0 * inch, 0.9 * inch],
                )
                t_style = [
                    ("BACKGROUND", (0, 0), (-1, 0), c_primary),
                    ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                    ("BOX", (0, 0), (-1, -1), 1, c_border),
                    ("INNERGRID", (0, 0), (-1, -1), 0.5, c_border),
                    ("TOPPADDING", (0, 0), (-1, -1), 5),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
                    ("LEFTPADDING", (0, 0), (-1, -1), 6),
                    ("RIGHTPADDING", (0, 0), (-1, -1), 6),
                ]
                # Alternating row background
                for r in range(1, len(task_rows)):
                    if r % 2 == 0:
                        t_style.append(("BACKGROUND", (0, r), (-1, r), c_light))

                task_table.setStyle(TableStyle(t_style))
                story.append(task_table)
                story.append(Spacer(1, 10))

        # 5. Full Transcript with Timestamps (if enabled)
        if options.get("include_transcript", True):
            transcript_text = note_data.get("transcript", "").strip()
            if transcript_text:
                story.append(Paragraph("Complete Audio Transcript", h2_style))
                story.append(
                    HRFlowable(
                        width="100%",
                        thickness=0.5,
                        color=c_border,
                        spaceBefore=0,
                        spaceAfter=8,
                    )
                )

                # Process transcript lines
                lines = transcript_text.split("\n")
                for line in lines:
                    line_str = line.strip()
                    if not line_str:
                        story.append(Spacer(1, 4))
                        continue

                    # If line has timestamp like [00:01:23] or Speaker:
                    if line_str.startswith("[") and "]" in line_str:
                        idx = line_str.find("]")
                        timestamp = line_str[: idx + 1]
                        rest = line_str[idx + 1 :].strip()
                        formatted_line = f"<font color='{c_accent.hexval()}'><b>{timestamp}</b></font> {rest}"
                        story.append(Paragraph(formatted_line, body_style))
                    else:
                        story.append(Paragraph(line_str, body_style))

        # Build Document
        doc.build(story)
        return output_path

    def export_docx(
        self,
        note_data: Dict[str, Any],
        output_path: str,
        options: Dict[str, bool],
    ) -> str:
        """Export note to a formatted Microsoft Word DOCX document using python-docx."""
        try:
            import docx
            from docx import Document
            from docx.shared import Inches, Pt, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.enum.table import WD_TABLE_ALIGNMENT
            from docx.oxml import parse_xml, OxmlElement
            from docx.oxml.ns import nsdecls, qn
        except ImportError as e:
            raise ImportError(
                f"python-docx is required for DOCX export: {e}. Install with 'pip install python-docx'."
            )

        doc = Document()

        # Set page margins
        for section in doc.sections:
            section.top_margin = Inches(0.75)
            section.bottom_margin = Inches(0.75)
            section.left_margin = Inches(0.75)
            section.right_margin = Inches(0.75)

        # Brand Colors
        c_primary_rgb = RGBColor(0x1E, 0x2B, 0x4B)  # #1E2B4B
        c_accent_rgb = RGBColor(0x6D, 0x59, 0xA7)  # #6D59A7
        c_sub_rgb = RGBColor(0x71, 0x80, 0x96)  # #718096

        # Helper: set background color of a table cell
        def set_cell_background(cell, hex_color: str):
            shading_elm = parse_xml(
                f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>'
            )
            cell._tc.get_or_add_tcPr().append(shading_elm)

        # Helper: set table borders
        def set_table_borders(table, hex_color: str = "E2DDD3"):
            tblPr = table._tbl.tblPr
            borders = parse_xml(
                f'<w:tblBorders {nsdecls("w")}>'
                f'<w:top w:val="single" w:sz="4" w:space="0" w:color="{hex_color}"/>'
                f'<w:bottom w:val="single" w:sz="4" w:space="0" w:color="{hex_color}"/>'
                f'<w:left w:val="single" w:sz="4" w:space="0" w:color="{hex_color}"/>'
                f'<w:right w:val="single" w:sz="4" w:space="0" w:color="{hex_color}"/>'
                f'<w:insideH w:val="single" w:sz="4" w:space="0" w:color="{hex_color}"/>'
                f'<w:insideV w:val="single" w:sz="4" w:space="0" w:color="{hex_color}"/>'
                f'</w:tblBorders>'
            )
            tblPr.append(borders)

        # 1. Header Banner
        p_sub = doc.add_paragraph()
        p_sub.paragraph_format.space_after = Pt(2)
        run_sub = p_sub.add_run("VOICENOTE AI • DOCUMENT EXPORT")
        run_sub.font.name = "Arial"
        run_sub.font.size = Pt(9)
        run_sub.font.color.rgb = c_sub_rgb
        run_sub.font.bold = True

        p_title = doc.add_paragraph()
        p_title.paragraph_format.space_after = Pt(12)
        run_title = p_title.add_run(note_data.get("title", "Voice Note"))
        run_title.font.name = "Arial"
        run_title.font.size = Pt(20)
        run_title.font.bold = True
        run_title.font.color.rgb = c_primary_rgb

        # 2. Metadata Table (if enabled)
        if options.get("include_metadata", True):
            tags_str = ", ".join(
                [f"#{t.lstrip('#')}" for t in note_data.get("tags", [])]
            ) or "General"

            meta_table = doc.add_table(rows=2, cols=4)
            meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER
            set_table_borders(meta_table, "E2DDD3")

            meta_entries = [
                ("Created Date:", str(note_data.get("created_at", "")), "Audio Duration:", str(note_data.get("duration", "00:00"))),
                ("Tags / Topics:", tags_str, "Generated By:", "VoiceNote AI Desktop"),
            ]

            for row_idx, row_data in enumerate(meta_entries):
                for col_idx in range(4):
                    cell = meta_table.cell(row_idx, col_idx)
                    set_cell_background(cell, "F8F7F4")
                    p = cell.paragraphs[0]
                    p.paragraph_format.space_after = Pt(2)
                    p.paragraph_format.space_before = Pt(2)
                    run = p.add_run(row_data[col_idx])
                    run.font.name = "Arial"
                    run.font.size = Pt(9)
                    if col_idx % 2 == 0:
                        run.font.bold = True
                        run.font.color.rgb = c_primary_rgb

            doc.add_paragraph()

        # 3. AI Executive Summary (if enabled)
        if options.get("include_summary", True):
            summary_text = note_data.get("summary")
            key_points = note_data.get("key_points", [])

            if summary_text or key_points:
                h_sum = doc.add_paragraph()
                h_sum.paragraph_format.space_before = Pt(12)
                h_sum.paragraph_format.space_after = Pt(4)
                r_sum = h_sum.add_run("AI Executive Summary")
                r_sum.font.name = "Arial"
                r_sum.font.size = Pt(13)
                r_sum.font.bold = True
                r_sum.font.color.rgb = c_primary_rgb

                if summary_text:
                    callout_table = doc.add_table(rows=1, cols=1)
                    callout_table.alignment = WD_TABLE_ALIGNMENT.CENTER
                    cell = callout_table.cell(0, 0)
                    set_cell_background(cell, "F8F7F4")
                    set_table_borders(callout_table, "6D59A7")
                    p_box = cell.paragraphs[0]
                    p_box.paragraph_format.space_before = Pt(6)
                    p_box.paragraph_format.space_after = Pt(6)
                    r_box = p_box.add_run(summary_text)
                    r_box.font.name = "Arial"
                    r_box.font.size = Pt(10)
                    r_box.font.italic = True
                    r_box.font.color.rgb = c_primary_rgb

                    doc.add_paragraph()

                if key_points:
                    p_kp_title = doc.add_paragraph()
                    p_kp_title.paragraph_format.space_before = Pt(6)
                    p_kp_title.paragraph_format.space_after = Pt(2)
                    r_kp_title = p_kp_title.add_run("Key Takeaways:")
                    r_kp_title.font.name = "Arial"
                    r_kp_title.font.size = Pt(10)
                    r_kp_title.font.bold = True

                    for kp in key_points:
                        p_bullet = doc.add_paragraph(style="List Bullet")
                        p_bullet.paragraph_format.space_after = Pt(2)
                        r_bp = p_bullet.add_run(kp)
                        r_bp.font.name = "Arial"
                        r_bp.font.size = Pt(10)

        # 4. Action Items & Extracted Tasks (if enabled)
        if options.get("include_tasks", True):
            tasks = note_data.get("tasks", [])
            if tasks:
                h_task = doc.add_paragraph()
                h_task.paragraph_format.space_before = Pt(14)
                h_task.paragraph_format.space_after = Pt(6)
                r_task = h_task.add_run("Extracted Action Items & Tasks")
                r_task.font.name = "Arial"
                r_task.font.size = Pt(13)
                r_task.font.bold = True
                r_task.font.color.rgb = c_primary_rgb

                task_table = doc.add_table(rows=len(tasks) + 1, cols=5)
                task_table.alignment = WD_TABLE_ALIGNMENT.CENTER
                set_table_borders(task_table, "E2DDD3")

                headers = ["Task Description", "Priority", "Assignee", "Due Date", "Status"]
                for c_idx, h_text in enumerate(headers):
                    cell = task_table.cell(0, c_idx)
                    set_cell_background(cell, "1E2B4B")
                    p = cell.paragraphs[0]
                    p.paragraph_format.space_before = Pt(4)
                    p.paragraph_format.space_after = Pt(4)
                    r = p.add_run(h_text)
                    r.font.name = "Arial"
                    r.font.size = Pt(9)
                    r.font.bold = True
                    r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

                for r_idx, t in enumerate(tasks, start=1):
                    row_data = [
                        t.get("title", ""),
                        t.get("priority", "Medium"),
                        t.get("assignee", "Unassigned"),
                        t.get("due_date", "TBD"),
                        t.get("status", "Pending"),
                    ]
                    bg_color = "F8F7F4" if r_idx % 2 == 0 else "FFFFFF"
                    for c_idx, val in enumerate(row_data):
                        cell = task_table.cell(r_idx, c_idx)
                        set_cell_background(cell, bg_color)
                        p = cell.paragraphs[0]
                        p.paragraph_format.space_before = Pt(3)
                        p.paragraph_format.space_after = Pt(3)
                        r = p.add_run(val)
                        r.font.name = "Arial"
                        r.font.size = Pt(9)
                        if c_idx == 1:
                            r.font.bold = True

                doc.add_paragraph()

        # 5. Full Audio Transcript (if enabled)
        if options.get("include_transcript", True):
            transcript_text = note_data.get("transcript", "").strip()
            if transcript_text:
                h_tr = doc.add_paragraph()
                h_tr.paragraph_format.space_before = Pt(14)
                h_tr.paragraph_format.space_after = Pt(6)
                r_tr = h_tr.add_run("Complete Audio Transcript")
                r_tr.font.name = "Arial"
                r_tr.font.size = Pt(13)
                r_tr.font.bold = True
                r_tr.font.color.rgb = c_primary_rgb

                lines = transcript_text.split("\n")
                for line in lines:
                    line_str = line.strip()
                    if not line_str:
                        continue
                    p_line = doc.add_paragraph()
                    p_line.paragraph_format.space_after = Pt(4)

                    if line_str.startswith("[") and "]" in line_str:
                        idx = line_str.find("]")
                        timestamp = line_str[: idx + 1]
                        rest = line_str[idx + 1 :].strip()

                        r_time = p_line.add_run(f"{timestamp} ")
                        r_time.font.name = "Arial"
                        r_time.font.size = Pt(9.5)
                        r_time.font.bold = True
                        r_time.font.color.rgb = c_accent_rgb

                        r_rest = p_line.add_run(rest)
                        r_rest.font.name = "Arial"
                        r_rest.font.size = Pt(9.5)
                    else:
                        r_plain = p_line.add_run(line_str)
                        r_plain.font.name = "Arial"
                        r_plain.font.size = Pt(9.5)

        # Save document
        doc.save(output_path)
        return output_path

    def export_txt(
        self,
        note_data: Dict[str, Any],
        output_path: str,
        options: Dict[str, bool],
    ) -> str:
        """Export note to a clean, structured Plain Text / Markdown document."""
        out = []

        title = note_data.get("title", "Voice Note")
        out.append("=" * 60)
        out.append(f"  {title.upper()}")
        out.append("  VoiceNote AI Desktop — Exported Document")
        out.append("=" * 60)
        out.append("")

        # 1. Metadata
        if options.get("include_metadata", True):
            tags_str = ", ".join(
                [f"#{t.lstrip('#')}" for t in note_data.get("tags", [])]
            ) or "General"
            out.append("[ METADATA ]")
            out.append(f"  • Date Created : {note_data.get('created_at', 'N/A')}")
            out.append(f"  • Duration     : {note_data.get('duration', '00:00')}")
            out.append(f"  • Categories   : {tags_str}")
            out.append(f"  • Export Date  : {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
            out.append("-" * 60)
            out.append("")

        # 2. Executive Summary
        if options.get("include_summary", True):
            summary = note_data.get("summary")
            key_points = note_data.get("key_points", [])
            if summary or key_points:
                out.append("[ AI EXECUTIVE SUMMARY ]")
                if summary:
                    out.append(f"  {summary}")
                    out.append("")
                if key_points:
                    out.append("  Key Takeaways:")
                    for kp in key_points:
                        out.append(f"    * {kp}")
                    out.append("")
                out.append("-" * 60)
                out.append("")

        # 3. Action Items & Tasks
        if options.get("include_tasks", True):
            tasks = note_data.get("tasks", [])
            if tasks:
                out.append("[ EXTRACTED ACTION ITEMS & TASKS ]")
                for i, t in enumerate(tasks, start=1):
                    status = t.get("status", "Pending")
                    box = "[X]" if status.lower() == "completed" else "[ ]"
                    priority = t.get("priority", "Medium")
                    assignee = t.get("assignee", "Unassigned")
                    due = t.get("due_date", "TBD")
                    out.append(
                        f"  {box} {i}. {t.get('title', '')}"
                    )
                    out.append(
                        f"      Priority: {priority} | Assignee: {assignee} | Due: {due} | Status: {status}"
                    )
                out.append("-" * 60)
                out.append("")

        # 4. Transcript
        if options.get("include_transcript", True):
            transcript = note_data.get("transcript", "").strip()
            if transcript:
                out.append("[ COMPLETE AUDIO TRANSCRIPT ]")
                out.append("")
                for line in transcript.split("\n"):
                    out.append(f"  {line}")
                out.append("")
                out.append("=" * 60)

        content = "\n".join(out)
        with open(output_path, "w", encoding="utf-8") as f:
            f.write(content)

        return output_path
